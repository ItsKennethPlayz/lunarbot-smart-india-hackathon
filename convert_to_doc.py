#!/usr/bin/env python3
"""
Markdown to DOC Converter
Converts the Team Task Delegation Framework markdown to Microsoft Word format.
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import markdown2
import re


def create_custom_styles(document):
    """Create custom styles for the document"""
    styles = document.styles

    # Title style
    try:
        title_style = styles['Title']
    except KeyError:
        title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Inches(0.25)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading styles
    for i in range(1, 4):
        heading_name = f'Heading {i}'
        try:
            heading_style = styles[heading_name]
        except KeyError:
            heading_style = styles.add_style(heading_name, WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Inches(0.2)
            heading_style.font.color.rgb = RGBColor(0, 51, 102)
        elif i == 2:
            heading_style.font.size = Inches(0.18)
            heading_style.font.color.rgb = RGBColor(51, 51, 51)
        else:
            heading_style.font.size = Inches(0.16)
            heading_style.font.color.rgb = RGBColor(102, 102, 102)


def process_markdown_content(content):
    """Process markdown content and convert to structured data"""
    # Convert markdown to HTML first
    html = markdown2.markdown(content, extras=['fenced-code-blocks', 'tables'])

    # Split content into sections based on headers
    sections = []
    current_section = {'level': 0, 'title': '', 'content': []}

    lines = content.split('\n')
    for line in lines:
        if line.strip():
            # Check for headers
            if line.startswith('#'):
                # Save previous section
                if current_section['title'] or current_section['content']:
                    sections.append(current_section)

                # Start new section
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                current_section = {'level': level, 'title': title, 'content': []}
            else:
                current_section['content'].append(line)

    # Don't forget the last section
    if current_section['title'] or current_section['content']:
        sections.append(current_section)

    return sections


def add_section_to_doc(doc, section):
    """Add a section to the document with appropriate styling"""
    if section['title']:
        # Add heading
        level = min(section['level'], 3)  # Limit to 3 levels
        if level == 1:
            heading = doc.add_heading(section['title'], level=1)
        elif level == 2:
            heading = doc.add_heading(section['title'], level=2)
        else:
            heading = doc.add_heading(section['title'], level=3)

    # Process content
    current_paragraph = None
    in_list = False

    for line in section['content']:
        line = line.strip()
        if not line:
            if current_paragraph:
                current_paragraph = None
            continue

        # Handle bullet points
        if line.startswith('- ') or line.startswith('* '):
            content = line[2:].strip()
            # Remove markdown formatting
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Bold
            content = re.sub(r'\*(.*?)\*', r'\1', content)      # Italic
            content = re.sub(r'`(.*?)`', r'\1', content)        # Code

            p = doc.add_paragraph(content, style='List Bullet')
            in_list = True

        # Handle bold headings (like **Primary Responsibilities:**)
        elif line.startswith('**') and line.endswith('**'):
            heading_text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            current_paragraph = None
            in_list = False

        # Handle regular text
        elif line and not line.startswith('#'):
            # Remove markdown formatting
            processed_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
            processed_line = re.sub(r'\*(.*?)\*', r'\1', processed_line)      # Italic
            processed_line = re.sub(r'`(.*?)`', r'\1', processed_line)        # Code

            if current_paragraph is None:
                current_paragraph = doc.add_paragraph(processed_line)
            else:
                current_paragraph.add_run(' ' + processed_line)
            in_list = False


def convert_markdown_to_doc(input_file, output_file):
    """Convert markdown file to DOC format"""
    try:
        # Read markdown file
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Create new document
        doc = Document()

        # Create custom styles
        create_custom_styles(doc)

        # Add title
        title = doc.add_heading('Smart India Hackathon: Autonomous Lunar Habitat Robot', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading('Team Task Delegation Framework (16-Week Research Program)', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Process content
        sections = process_markdown_content(content)

        for section in sections:
            add_section_to_doc(doc, section)

        # Save document
        doc.save(output_file)
        print(f"✅ Successfully converted {input_file} to {output_file}")
        return True

    except Exception as e:
        print(f"❌ Error converting to DOC: {str(e)}")
        return False


def main():
    """Main function"""
    # Set file paths
    script_dir = Path(__file__).parent
    input_file = script_dir / "Team_Task_Delegation_Framework.md"
    output_file = script_dir / "Team_Task_Delegation_Framework.docx"

    # Check if input file exists
    if not input_file.exists():
        print(f"❌ Input file not found: {input_file}")
        print("Available files:")
        for file in script_dir.glob("*.md"):
            print(f"  - {file.name}")
        return False

    print(f"🚀 Converting {input_file.name} to DOC format...")
    print(f"📁 Input: {input_file}")
    print(f"📄 Output: {output_file}")

    success = convert_markdown_to_doc(input_file, output_file)

    if success:
        print(f"\n🎉 Conversion completed successfully!")
        print(f"📋 DOC file created: {output_file}")
        print(f"📊 File size: {output_file.stat().st_size / 1024:.1f} KB")
    else:
        print(f"\n💥 Conversion failed!")
        return False

    return True


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
